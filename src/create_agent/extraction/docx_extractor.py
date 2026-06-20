"""Extract text from .docx files using python-docx."""

from __future__ import annotations

from pathlib import Path


def extract_docx_text(file_path: str | Path) -> str:
    """Extract all text from a .docx file.

    Reads paragraphs and table cells, joining them with newlines.

    Args:
        file_path: Path to the .docx file.

    Returns:
        Extracted text content.

    Raises:
        ImportError: If python-docx is not installed.
        ValueError: If the file is not a valid .docx file.
    """
    from docx import Document

    path = Path(file_path)
    if not path.suffix.lower() in (".docx", ".docm"):
        raise ValueError(f"Not a .docx file: {path}")

    doc = Document(str(path))

    parts: list[str] = []

    # Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Extract table content
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(row_texts))

    return "\n".join(parts)
